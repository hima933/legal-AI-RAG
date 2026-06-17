"""
Multi-format document processing
Supports: PDF, DOCX, TXT, RTF, Images (OCR), DOC
"""

import os
import logging
from typing import List, Tuple, Optional, Dict, Any
from pathlib import Path
from pypdf import PdfReader
import mimetypes

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Base processor for documents"""
    
    def __init__(self):
        self.supported_formats = {}
    
    def process(self, file_path: str) -> Tuple[str, List[str]]:
        """
        Process document and extract text
        Returns: (full_text, extracted_metadata_list)
        """
        raise NotImplementedError


class PDFProcessor(DocumentProcessor):
    """Process PDF documents"""
    
    def __init__(self):
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf required. Install: pip install pypdf")
        
        self.PdfReader = PdfReader
    
    def process(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text and metadata from PDF"""
        try:
            full_text = []
            metadata = []
            
            with open(file_path, "rb") as f:
                reader = self.PdfReader(f)
                
                # Extract metadata
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        metadata.append(f"{key}: {value}")
                
                # Extract text from all pages
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text.append(f"[Page {page_num}]\n{text}")
            
            return "\n".join(full_text), metadata
        
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise


class DOCXProcessor(DocumentProcessor):
    """Process DOCX (Microsoft Word) documents"""
    
    def __init__(self):
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required. Install: pip install python-docx")
        
        self.Document = Document
    
    def process(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text and metadata from DOCX"""
        try:
            doc = self.Document(file_path)
            
            # Extract metadata
            metadata = []
            core_props = doc.core_properties
            if core_props.title:
                metadata.append(f"Title: {core_props.title}")
            if core_props.author:
                metadata.append(f"Author: {core_props.author}")
            if core_props.created:
                metadata.append(f"Created: {core_props.created}")
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        full_text.append(row_text)
            
            return "\n".join(full_text), metadata
        
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise


class TXTProcessor(DocumentProcessor):
    """Process plain text documents"""
    
    def process(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from TXT files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            metadata = [
                f"File: {Path(file_path).name}",
                f"Size: {os.path.getsize(file_path)} bytes"
            ]
            
            return text, metadata
        
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            # Try different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    text = f.read()
                return text, metadata
            except:
                raise


class RTFProcessor(DocumentProcessor):
    """Process RTF (Rich Text Format) documents"""
    
    def __init__(self):
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf required. Install: pip install striprtf")
        
        self.rtf_to_text = rtf_to_text
    
    def process(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from RTF"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                rtf_content = f.read()
            
            text = self.rtf_to_text(rtf_content)
            metadata = [f"File: {Path(file_path).name}"]
            
            return text, metadata
        
        except Exception as e:
            logger.error(f"Error processing RTF: {e}")
            raise


class ImageProcessor(DocumentProcessor):
    """Process images with OCR (Optical Character Recognition)"""
    
    def __init__(self):
        try:
            from PIL import Image
        except ImportError:
            raise ImportError("Pillow required. Install: pip install Pillow")
        
        try:
            import pytesseract
        except ImportError:
            raise ImportError("pytesseract required. Install: pip install pytesseract")
        
        self.Image = Image
        self.pytesseract = pytesseract
    
    def process(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from images using OCR"""
        try:
            image = self.Image.open(file_path)
            
            # Extract text via OCR
            text = self.pytesseract.image_to_string(image)
            
            metadata = [
                f"File: {Path(file_path).name}",
                f"Format: {image.format}",
                f"Size: {image.size}",
                f"Mode: {image.mode}"
            ]
            
            return text, metadata
        
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise


class DocumentFactory:
    """Factory for creating appropriate document processors"""
    
    _processors = {
        "pdf": PDFProcessor,
        "docx": DOCXProcessor,
        "doc": DOCXProcessor,  # Can use docx processor
        "txt": TXTProcessor,
        "rtf": RTFProcessor,
        "jpg": ImageProcessor,
        "jpeg": ImageProcessor,
        "png": ImageProcessor,
        "bmp": ImageProcessor,
        "tiff": ImageProcessor,
    }
    
    @staticmethod
    def get_processor(file_path: str) -> DocumentProcessor:
        """Get appropriate processor for file type"""
        file_ext = Path(file_path).suffix.lower().lstrip(".")
        
        if file_ext not in DocumentFactory._processors:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        processor_class = DocumentFactory._processors[file_ext]
        return processor_class()
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(file_path).suffix.lower().lstrip(".")
        return file_ext in DocumentFactory._processors


class TextChunker:
    """Split documents into chunks for embedding"""
    
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size  # words
        self.overlap = overlap  # words
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks
        Better for semantic search
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk = " ".join(words[i:i + self.chunk_size])
            if len(chunk.split()) >= 50:  # Minimum chunk size
                chunks.append(chunk)
        
        return chunks
    
    def chunk_by_sections(self, text: str) -> List[str]:
        """
        Split by legal sections/clauses
        Better for legal documents
        """
        # Split by common legal patterns
        import re
        
        patterns = [
            r"^section\s+\d+",
            r"^article\s+\d+",
            r"^clause\s+\d+",
            r"^§\s*\d+",
        ]
        
        chunks = []
        current_chunk = ""
        
        for line in text.split("\n"):
            if any(re.match(pattern, line.lower()) for pattern in patterns):
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = line
            else:
                current_chunk += "\n" + line
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks


class DocumentManager:
    """
    High-level interface for document processing
    Handles multi-format documents with validation
    """
    
    def __init__(self):
        self.chunker = TextChunker()
    
    def process_file(self, file_path: str, chunk: bool = True) -> Tuple[str, List[str], Dict[str, Any]]:
        """
        Process single document
        Returns: (full_text, chunks, metadata)
        """
        
        # Validate file
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > 50:  # 50 MB limit
            raise ValueError(f"File too large: {file_size_mb:.1f}MB (max 50MB)")
        
        # Check if supported
        if not DocumentFactory.is_supported(file_path):
            raise ValueError(f"Unsupported file type: {Path(file_path).suffix}")
        
        # Process
        processor = DocumentFactory.get_processor(file_path)
        logger.info(f"Processing {Path(file_path).name} with {processor.__class__.__name__}")
        
        full_text, raw_metadata = processor.process(file_path)
        
        # Chunk if requested
        if chunk:
            file_ext = Path(file_path).suffix.lower().lstrip(".")
            section_chunks: List[str] = []
            if file_ext in {"pdf", "doc", "docx", "rtf", "txt"}:
                section_chunks = self.chunker.chunk_by_sections(full_text)
                section_chunks = [c for c in section_chunks if len(c.split()) >= 40]

            if len(section_chunks) >= 3:
                chunks = section_chunks
            else:
                chunks = self.chunker.chunk_text(full_text)
        else:
            chunks = [full_text]
        
        logger.info(f"Extracted {len(chunks)} chunks from {Path(file_path).name}")

        metadata = {
            "source_file": Path(file_path).name,
            "file_size_mb": round(file_size_mb, 3),
            "processor": processor.__class__.__name__,
            "metadata_items": raw_metadata,
            "chunk_count": len(chunks),
            "char_count": len(full_text),
            "word_count": len(full_text.split()),
        }

        return full_text, chunks, metadata
    
    def process_multiple_files(self, file_paths: List[str]) -> Tuple[List[str], List[Dict[str, Any]]]:
        """
        Process multiple documents
        Returns: (all_chunks, all_metadata)
        """
        all_chunks = []
        all_metadata = []
        
        for file_path in file_paths:
            try:
                _, chunks, metadata = self.process_file(file_path)
                all_chunks.extend(chunks)
                all_metadata.append(metadata)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                all_metadata.append({
                    "source_file": Path(file_path).name,
                    "error": str(e)
                })
        
        return all_chunks, all_metadata
